import re
from docx.shared import Mm, Pt
from checker.utils import map_alignment, mm_to_pt, clean_empty_paragraphs, insert_page_breaks

def before_check(doc):
    clean_empty_paragraphs(doc, heading_levels=(1, 2))
    insert_page_breaks(doc)

def check_and_fix_style(doc, config):
    issues = []
    skipping_title_list = True
    before_check(doc)    

    # 1. Проверка и установка полей страницы
    margins_cfg = config.get('margins', {})
    for section in doc.sections:
        # Текущие поля (можно добавить проверку несоответствий)
        # Устанавливаем нужные поля
        section.left_margin = Mm(margins_cfg.get('left', 20))
        section.right_margin = Mm(margins_cfg.get('right', 20))
        section.top_margin = Mm(margins_cfg.get('top', 20))
        section.bottom_margin = Mm(margins_cfg.get('bottom', 20))

    # 2. Проход по параграфам
    for paragraph in doc.paragraphs:
        
        if skipping_title_list:
            if "Минск 2025" in paragraph.text:
                skipping_title_list = False
            continue
        
        text = paragraph.text.strip()
        if not text:
            continue

        style_name = paragraph.style.name
        # Если это заголовок
        match = re.match(r'Heading (\d+)', style_name)
        if match:
            level = int(match.group(1))
            issues += _fix_heading(paragraph, config.get('headings', {}), level)
        else:
            # Обычный параграф (Normal или иной)
            issues += _fix_paragraph(paragraph, config.get('font', {}), config.get('paragraph', {}))

    # 3. Проверка таблиц
    for table in doc.tables:
        issues += _fix_table(table, config.get('tables', {}))

    # 4. Проверка подписей к изображениям
    # Подписи обычно хранятся в подпараграфе под рисунком, проверим тексты, начинающиеся с "Рисунок"
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith('Рисунок'):
            issues += _fix_image_caption(paragraph, config.get('images', {}))

    return issues


def _fix_paragraph(paragraph, font_cfg, parag_cfg):
    issues = []

    # Шрифт
    for run in paragraph.runs:
        if run.font.name != font_cfg.get('name'):
            run.font.name = font_cfg.get('name')
            issues.append(f"Изменён шрифт параграфа: '{paragraph.text[:30]}...' -> {font_cfg.get('name')}")
        if run.font.size and run.font.size.pt != font_cfg.get('size'):
            run.font.size = Pt(font_cfg.get('size'))
            issues.append(f"Изменён размер шрифта: '{paragraph.text[:30]}...' -> {font_cfg.get('size')} pt")
        if font_cfg.get('bold', False) and not run.font.bold:
            run.font.bold = False
            issues.append(f"Отменён полужирный стиль: '{paragraph.text[:30]}...'")
        if font_cfg.get('italic', False) and not run.font.italic:
            run.font.italic = False
            issues.append(f"Отменён курсив: '{paragraph.text[:30]}...'")
        if font_cfg.get('underlined', False) and not run.font.underline:
            run.font.underline = False
            issues.append(f"Отменено подчёркивание: '{paragraph.text[:30]}...'")

    # Абзац
    # Отступ первой строки
    desired_indent = mm_to_pt(parag_cfg.get('first_line_indent_mm', 0))
    current_indent = paragraph.paragraph_format.first_line_indent.pt if paragraph.paragraph_format.first_line_indent else 0
    if abs(current_indent - desired_indent) > 0.1:
        paragraph.paragraph_format.first_line_indent = Pt(desired_indent)
        issues.append(f"Изменён отступ первой строки: '{paragraph.text[:30]}...' -> {parag_cfg.get('first_line_indent_mm')} мм")

    # Межстрочный интервал
    current_spacing = paragraph.paragraph_format.line_spacing if paragraph.paragraph_format.line_spacing else 1.0
    if abs(current_spacing - parag_cfg.get('line_spacing', 1.0)) > 0.01:
        paragraph.paragraph_format.line_spacing = parag_cfg.get('line_spacing')
        issues.append(f"Изменён межстрочный интервал: '{paragraph.text[:30]}...' -> {parag_cfg.get('line_spacing')}")

    # Выравнивание
    desired_align = map_alignment(parag_cfg.get('alignment', 'justify'))
    if paragraph.alignment != desired_align:
        paragraph.alignment = desired_align
        issues.append(f"Изменено выравнивание: '{paragraph.text[:30]}...' -> {parag_cfg.get('alignment')}")

    return issues

def _fix_heading(paragraph, headings_cfg, level):
    issues = []
    key = f"level_{level}"
    if key in headings_cfg:
        cfg = headings_cfg[key]
        for run in paragraph.runs:
            # Шрифт заголовка всегда тот же, но размер может отличаться
            if run.font.name != cfg.get('font_name', 'Times New Roman'):
                run.font.name = cfg.get('font_name', 'Times New Roman')
                issues.append(f"Изменён шрифт заголовка уровня {level}: '{paragraph.text[:30]}...' -> {cfg.get('font_name')}")
            if run.font.size and run.font.size.pt != cfg.get('font_size'):
                run.font.size = Pt(cfg.get('font_size'))
                issues.append(f"Изменён размер шрифта заголовка уровня {level}: '{paragraph.text[:30]}...' -> {cfg.get('font_size')} pt")
            if cfg.get('bold', False) and not run.font.bold:
                run.font.bold = True
                issues.append(f"Применён полужирный стиль к заголовку уровня {level}: '{paragraph.text[:30]}...'")
            if cfg.get('italic', False) and not run.font.italic:
                run.font.italic = True
                issues.append(f"Применён курсив к заголовку уровня {level}: '{paragraph.text[:30]}...'")
            # Выравнивание заголовка
            desired_align = map_alignment(cfg.get('alignment', 'left'))
            if paragraph.alignment != desired_align:
                paragraph.alignment = desired_align
                issues.append(f"Изменено выравнивание заголовка уровня {level}: '{paragraph.text[:30]}...' -> {cfg.get('alignment')}")
            # Отступ перед заголовком (space_before)
            current_space_before = paragraph.paragraph_format.space_before.pt if paragraph.paragraph_format.space_before else 0
            desired_space_before = cfg.get('space_before', 0)
            if abs(current_space_before - desired_space_before) > 0.1:
                paragraph.paragraph_format.space_before = Pt(desired_space_before)
                issues.append(f"Изменён отступ после заголовка уровня {level}: '{paragraph.text[:30]}...' -> {desired_space_before} pt")
            # Отступ после заголовка (space_after)
            current_space_after = paragraph.paragraph_format.space_after.pt if paragraph.paragraph_format.space_after else 0
            desired_space_after = cfg.get('space_after', 0)
            if abs(current_space_after - desired_space_after) > 0.1:
                paragraph.paragraph_format.space_after = Pt(desired_space_after)
                issues.append(f"Изменён отступ после заголовка уровня {level}: '{paragraph.text[:30]}...' -> {desired_space_after} pt")
    return issues


def _fix_table(table, tables_cfg):
    issues = []
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if run.font.name != tables_cfg.get('font_name'):
                        run.font.name = tables_cfg.get('font_name')
                        issues.append(f"Изменён шрифт в таблице: '{cell.text[:30]}...' -> {tables_cfg.get('font_name')}")
                    if run.font.size and run.font.size.pt != tables_cfg.get('font_size'):
                        run.font.size = Pt(tables_cfg.get('font_size'))
                        issues.append(f"Изменён размер шрифта в таблице: '{cell.text[:30]}...' -> {tables_cfg.get('font_size')} pt")
            # Межстрочный в ячейке
            current_spacing = paragraph.paragraph_format.line_spacing if paragraph.paragraph_format.line_spacing else 1.0
            if abs(current_spacing - tables_cfg.get('line_spacing', 1.0)) > 0.01:
                paragraph.paragraph_format.line_spacing = tables_cfg.get('line_spacing')
                issues.append(f"Изменён межстрочный интервал в таблице: '{cell.text[:30]}...' -> {tables_cfg.get('line_spacing')}")
            desired_align = map_alignment(tables_cfg.get('alignment', 'center'))
            if paragraph.alignment != desired_align:
                paragraph.alignment = desired_align
                issues.append(f"Изменено выравнивание подписи изображения: '{paragraph.text[:30]}...' -> {tables_cfg.get('alignment')}")
    return issues

def _fix_image_caption(paragraph, images_cfg):
    issues = []
    # Устанавливаем нужный шрифт и размер для подписи
    for run in paragraph.runs:
        if run.font.name != images_cfg.get('caption_font_name'):
            run.font.name = images_cfg.get('caption_font_name')
            issues.append(f"Изменён шрифт подписи изображения: '{paragraph.text[:30]}...' -> {images_cfg.get('caption_font_name')}")
        if run.font.size and run.font.size.pt != images_cfg.get('caption_font_size'):
            run.font.size = Pt(images_cfg.get('caption_font_size'))
            issues.append(f"Изменён размер шрифта подписи изображения: '{paragraph.text[:30]}...' -> {images_cfg.get('caption_font_size')} pt")
    desired_align = map_alignment(images_cfg.get('alignment', 'center'))
    if paragraph.alignment != desired_align:
        paragraph.alignment = desired_align
        issues.append(f"Изменено выравнивание подписи изображения: '{paragraph.text[:30]}...' -> {images_cfg.get('alignment')}")
    return issues