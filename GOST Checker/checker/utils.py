import re
import requests
import os
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from datetime import datetime

def map_alignment(align_str):
    """
    Преобразует строку 'left', 'right', 'center', 'justify' в соответствующее значение WD_PARAGRAPH_ALIGNMENT.
    """
    align_str = align_str.lower()
    if align_str == 'left':
        return WD_PARAGRAPH_ALIGNMENT.LEFT
    if align_str == 'right':
        return WD_PARAGRAPH_ALIGNMENT.RIGHT
    if align_str == 'center':
        return WD_PARAGRAPH_ALIGNMENT.CENTER
    return WD_PARAGRAPH_ALIGNMENT.JUSTIFY

def mm_to_pt(mm_value):
    """
    Конвертирует миллиметры (mm) в пункты (pt).
    В одном пункте (pt) ≈ 0.352777778 мм.
    """
    return mm_value / 0.352777778

def has_page_break(paragraph):
    for run in paragraph.runs:
        for br in run._element.findall('.//w:br', namespaces=run._element.nsmap):
            if br.get(qn('w:type')) == 'page':
                return True
    return False

def insert_page_breaks(doc):
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.style.name == 'Heading 1':
            # Пропускаем первый параграф
            if i == 0:
                continue

            prev_para = doc.paragraphs[i - 1]
            if not has_page_break(prev_para):
                new_para = paragraph.insert_paragraph_before()
                new_para.add_run().add_break(WD_BREAK.PAGE)


def is_empty_paragraph(paragraph):
    """Проверяет, что абзац пустой — нет текста и runs пустые."""
    return not paragraph.text.strip() and all(not run.text.strip() for run in paragraph.runs)

def clean_empty_paragraphs(doc, heading_levels=(1, 2)):
    skipping_title_list = True
    allowed_styles = [f'Heading {lvl}' for lvl in heading_levels] + [f'Заголовок {lvl}' for lvl in heading_levels]
    i = 1
    for paragraph in doc.paragraphs:
        if skipping_title_list:
            i+=1
            if "Минск 2025" in paragraph.text:
                skipping_title_list = False
            continue
    while i < len(doc.paragraphs):
        paragraph = doc.paragraphs[i]
        if paragraph.style.name in allowed_styles:
            prev_para = doc.paragraphs[i - 1]
            if is_empty_paragraph(prev_para):
                p = prev_para._element
                p.getparent().remove(p)
                i -= 1  # после удаления список сдвинулся назад
        i += 1
        
        
def fetch_title(url):
    """
    Пытается получить название ресурса по URL:
    - Для PDF: если сервер отдает HTML (оболочка с заголовком) — пытаемся из HTML
    - Если действительно PDF (Content-Type application/pdf) — берём имя файла
    Общая логика:
    1. HEAD запрос: получить Content-Type
    2. Если text/html — парсим HTML по предыдущим правилам
    3. Если application/pdf или url оканчивается на .pdf — берём имя файла
    4. Иначе — парсим HTML
    """
    # Определяем датуум
    def title_from_html(html_text):
        soup = BeautifulSoup(html_text, 'html.parser')
        # 1. citation_title
        citation = soup.find('meta', attrs={'name': 'citation_title'})
        if citation and citation.get('content'):
            return citation['content'].strip()
        # 2. title
        if soup.title and soup.title.string:
            return soup.title.string.strip()
        # 3. og:title
        og = soup.find('meta', property='og:title')
        if og and og.get('content'):
            return og['content'].strip()
        # 4. meta name=title
        meta_title = soup.find('meta', attrs={'name': 'title'})
        if meta_title and meta_title.get('content'):
            return meta_title['content'].strip()
        # 5. h1
        h1 = soup.find('h1')
        if h1 and h1.get_text(strip=True):
            return h1.get_text(strip=True)
        # 6. h2
        h2 = soup.find('h2')
        if h2 and h2.get_text(strip=True):
            return h2.get_text(strip=True)
        # 7. meta description
        meta_desc = soup.find('meta', attrs={'name':'description'})
        if meta_desc and meta_desc.get('content'):
            return meta_desc['content'].strip()
        return None

    try:
        # Сначала HEAD
        head = requests.head(url, allow_redirects=True, timeout=5)
        content_type = head.headers.get('Content-Type', '')
        # Если PDF
        if 'application/pdf' in content_type or url.lower().endswith('.pdf'):
            # если сервер отдает HTML
            if 'text/html' in content_type:
                get = requests.get(url, timeout=5)
                res = title_from_html(get.text)
                if res:
                    return res
            # иначе возвращаем имя файла
            return os.path.splitext(os.path.basename(url.split('?')[0]))[0]
        # Для остальных, если HTML
        if 'text/html' in content_type:
            get = requests.get(url, timeout=5)
            res = title_from_html(get.text)
            if res:
                return res
    except Exception:
        pass
    # fallback: если html парсинг не помог
    # попытаться получить html
    try:
        resp = requests.get(url, timeout=5)
        resp.raise_for_status()
        res = title_from_html(resp.text)
        if res:
            return res
    except Exception:
        pass
    return url


def process_bibliography(doc: Document, bib_cfg: dict):
    """
    Находит раздел "СПИСОК ИСТОЧНИКОВ" и форматирует интернет-источники по шаблону из config.yaml.
    Поддерживает PDF-файлы и веб-страницы.
    """
    in_bib = False
    idx = 0
    today = datetime.today().strftime("%d.%m.%Y")
    template = bib_cfg.get(
        'format_template',
        '«{title}» [Электронный ресурс]. Режим доступа: {url} Дата доступа: {access_date}.'
    )

    for paragraph in doc.paragraphs:
        if not in_bib:
            if re.search(r'СПИСОК ИСТОЧНИКОВ', paragraph.text, re.IGNORECASE):
                in_bib = True
            continue

        text = paragraph.text.strip()
        if not text:
            continue

        match = re.search(r'(https?://[^\s]+)', text)
        if not match:
            continue

        url = match.group(1)
        idx += 1

        # Если ссылка на PDF, то берем имя файла без параметров
        if url.lower().endswith('.pdf'):
            title = os.path.splitext(os.path.basename(url.split('?')[0]))[0]
        else:
            title = fetch_title(url)

        new_text = template.format(
            index=idx,
            title=title,
            url=url,
            access_date=today
        )

        paragraph.clear()
        run = paragraph.add_run(new_text)
        run.font.name = bib_cfg.get('font_name', 'Times New Roman')
        run.font.size = Pt(bib_cfg.get('font_size', 14))
        paragraph.paragraph_format.space_after = Pt(bib_cfg.get('space_after', 6))
        paragraph.paragraph_format.line_spacing = bib_cfg.get('line_spacing', 1.0)
